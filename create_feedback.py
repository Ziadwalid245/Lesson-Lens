from docx import Document



def create_feedback_doc(feedback, output_path):
    doc = Document()
    doc.add_heading('Student Feedback', level=1)
    doc.add_heading('Lesson Summary', level=2)
    doc.add_paragraph(feedback.lesson_summary)
    vocab_table = doc.add_table(rows=1, cols=3)
    grammar_table = doc.add_table(rows=1, cols=3)
    corrections_table = doc.add_table(rows=1, cols=3)
    vocab_table.style = 'Table Grid'
    grammar_table.style = 'Table Grid'
    corrections_table.style = 'Table Grid'
   
    
    
    
    vocab_table_headers = vocab_table.rows[0].cells
    vocab_table_headers[0].text = 'Word'
    vocab_table_headers[1].text = 'Definition' 
    vocab_table_headers[2].text = 'Sentence in passage/Audio'

    for vocab in feedback.vocab_items:
        row_cells = vocab_table.add_row().cells
        row_cells[0].text = vocab.word
        row_cells[1].text = vocab.definition
        row_cells[2].text = vocab.in_context
    grammar_table_headers = grammar_table.rows[0].cells
    grammar_table_headers[0].text = 'Grammar Point'
    grammar_table_headers[1].text = 'Form'  
    grammar_table_headers[2].text = 'Usage'
    for gp in feedback.grammar_points:
        row_cells = grammar_table.add_row().cells
        row_cells[0].text = gp.explained_garmmar
        row_cells[1].text = gp.Form
        row_cells[2].text = gp.Usage
    corrections_table_headers = corrections_table.rows[0].cells
    corrections_table_headers[0].text = 'Sentence with Error'
    corrections_table_headers[1].text = 'Corrected Sentence'
    corrections_table_headers[2].text = 'Explanation'
    for correction in feedback.corrections:
        row_cells = corrections_table.add_row().cells
        row_cells[0].text = correction.original
        row_cells[1].text = correction.corrected
        row_cells[2].text = correction.explanation
    doc.add_heading('Positive Feedback', level=2)
    doc.add_paragraph(feedback.positive_feedback)
    doc.add_heading('Areas for Improvement', level=2)
    doc.add_paragraph(feedback.improvement_areas)
    doc.save(output_path)

